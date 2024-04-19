from selenium import webdriver
from send2trash import send2trash
from itertools import cycle
import os
from docx import Document
from docx.shared import Inches
import time
from datetime import datetime
import sys
from operaciones_mdp import OperacionesMDP
import xlwings as xw
import pyautogui
from selenium.common.exceptions import NoSuchElementException, TimeoutException, ElementNotInteractableException, StaleElementReferenceException, ElementClickInterceptedException
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By

from selenium.webdriver import Keys, ActionChains


class Motor:
    txt_username = '//*[@id="txtUser"]'
    txt_password = '//*[@id="txtPassword"]'
    txt_rut_ingresar = '//*[@id="ctl00_ContenedorPrimario_txtRut"]'
    txt_nombre_ingresar = '//*[@id="ctl00_ContenedorPrimario_txtNombre"]'
    txt_monto_ingresar = '//*[@id="ctl00_ContenedorPrimario_txtMonto"]'
    txt_numero_operacion = '//*[@id="ctl00_ContenedorPrimario_Tab_TabPanel3_txtNumeroOperacion"]'
    txt_cuenta_ingresar = '//*[@id="ctl00_ContenedorPrimario_txtfld_mdptrd_59"]'
    btn_entry_user = '//*[@id="Button1"]'
    btn_menu_transacciones = '//*[@id="ctl00_Accordion1"]/div[3]/a'
    btn_menu_ingresar = '//*[@id="ctl00_ctl08_btnIngres"]' 
    btn_menu_verificar = '//*[@id="ctl00_ctl08_btnVerifc"]'
    btn_menu_gestion = '//*[@id="ctl00_Accordion1"]/div[9]/a'
    btn_menu_consulta = '//*[@id="ctl00_ctl14_btnConsulta"]'
    btn_ingresar_nueva = '//*[@id="ctl00_ContenedorPrimario_btnNueva"]'
    btn_continuar_ingresar = '//*[@id="ctl00_ContenedorPrimario_btnContinuar"]'
    btn_confirmar_ingresar = '//*[@id="ctl00_ContenedorPrimario_btnBoton1"]'
    btn_beneficiario_ingresar = '//*[@id="ctl00_ContenedorPrimario_Tab_CliBanCob_TabPanel1_btnBeneficiario"]'
    btn_aceptar_ventana = '//*[@id="ctl00_ContenedorPrimario_btnAceptarMensaje"]'
    btn_home = '//*[@id="ctl00_btnHome"]'
    btn_numero_operacion = '//*[@id="__tab_ctl00_ContenedorPrimario_Tab_TabPanel3"]'
    btn_numero_operacion_buscar = '//*[@id="ctl00_ContenedorPrimario_Tab_TabPanel3_btnBuscar2"]'
    btn_cerrar_sesion = '//*[@id="ctl00_btnExit"]'
    btn_aceptar_cerrar_sesion = '//*[@id="ctl00_ButtonOk"]'
    btn_aprobar_verificar = '//*[@id="ctl00_ContenedorPrimario_btnAprobar"]'
    btn_grabar_beneficiario = '//*[@id="ctl00_ContenedorPrimario_btnGrabar"]'
    btn_error_usuario = '//*[@id="btnAceptarMensaje"]'
    casilla_numero_operacion = '//*[@id="ctl00_ContenedorPrimario_dgVerificar_ctl01_chkAll"]'
    slc_forma_pago = '//*[@id="ctl00_ContenedorPrimario_cmbFormaPago"]'
    slc_producto = '//*[@id="ctl00_ContenedorPrimario_cmbProductos"]'
    slc_banco = '//*[@id="ctl00_ContenedorPrimario_Tab_CliBanCob_TabPanel1_cmbBanco_Cliente"]'
    #//*[@id="ctl00_ContenedorPrimario_lblMensaje"] 
    
    def __init__(self,archivo_xlsx: str):
        self.app = xw.App(visible=False)
        self.archivo_xlsx = archivo_xlsx
        self.workbook = self.app.books.open(self.archivo_xlsx)
        self.hoja = self.workbook.sheets[0]
        
        options = webdriver.ChromeOptions()
        service = ChromeService(executable_path="./drivers/chromedriver.exe")
        self.driver = webdriver.Chrome(service=service, options=options)
        self.action = ActionChains(self.driver)
        self.document = Document()
        self.action_bot = OperacionesMDP(self.driver)
        self.driver.maximize_window()

    def ejecucion_motor(self):
        self.driver.get("http://...")
        self.document.add_heading('Reporte de Automatización', 0)
        range1 = self.hoja['A2:M2'].expand('down').rows
        self.e=1
        self.list_number = 'List Number'
        for i in range1:
            self.document.add_heading(f'Fila {self.e} de Excel', level=1)
            try:
                self.iniciar_sesion_ingresador(i) #INICIA USUARIO INGRESADOR
                time.sleep(5)
                if "Sistema " in self.driver.title:
                    self.action_bot.presionar_boton(self.btn_menu_transacciones)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_menu_ingresar)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_ingresar_nueva)
                    if i[5].value == 1:
                        self.ingresar_mt103(i)
                        time.sleep(10)
                    elif i[5].value == 2:
                        self.ingresar_mt202(i)
                        time.sleep(10)
                    else:
                        raise ValueError("No ha ingresado ningún tipo de operación\nPor favor revise el Excel y vuelva a intentar...") 
                    WebDriverWait(self.driver, timeout=90).until(EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_ContenedorPrimario_Tab_Panel1_txtfld_mdpope_sys_num"]')))           
                    elemento = self.driver.find_element(By.XPATH, '//*[@id="ctl00_ContenedorPrimario_Tab_Panel1_txtfld_mdpope_sys_num"]')
                    print("buscando elemento del numero..")
                    numero_operacion = elemento.get_attribute('value')
                    print("almacenando...")
                    i[10].value = numero_operacion
                    print(numero_operacion)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_confirmar_ingresar)
                    print(f"Confirmando operación {numero_operacion}...")
                    WebDriverWait(self.driver, timeout=60).until(EC.presence_of_element_located((By.XPATH, self.btn_aceptar_ventana)))
                    time.sleep(3)
                    img_confirmacion_ingreso = f"./Screenshots/{self.e}_MDP_CONFIRMACION_{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
                    screenshot1 = pyautogui.screenshot()
                    screenshot1.save(img_confirmacion_ingreso)
                    #Inserción en Word
                    self.document.add_paragraph(f'Confirmación de ingreso de operación {numero_operacion}', style=self.list_number)
                    self.document.add_picture(img_confirmacion_ingreso, width=Inches(6.73))
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_aceptar_ventana)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_home)
                    time.sleep(3)
                    self.consultar_operacion(numero_operacion)
                    time.sleep(10)
                    img_consulta_ingreso = f"./Screenshots/{self.e}_MDP_CONSULTA_{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
                    screenshot2 = pyautogui.screenshot()
                    screenshot2.save(img_consulta_ingreso)
                    time.sleep(3)
                    self.document.add_paragraph(f'Consulta de ingreso de operación {numero_operacion}', style=self.list_number)
                    time.sleep(3)
                    self.document.add_picture(img_consulta_ingreso, width=Inches(6.73))
                    time.sleep(5)
                    self.cerrar_sesion()
                    self.iniciar_sesion_verificador(i) #INICIA USUARIO VERIFICADOR
                    self.verificar_con_verificador(numero_operacion)
                    WebDriverWait(self.driver, timeout=60).until(EC.presence_of_element_located((By.XPATH, self.btn_aceptar_ventana)))
                    time.sleep(3)
                    img_verificacion_operacion = f"./Screenshots/{self.e}_MDP101_CONFIRMACION_VERIFICADO{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
                    screenshot3 = pyautogui.screenshot()
                    screenshot3.save(img_verificacion_operacion)
                    time.sleep(3)
                    self.document.add_paragraph(f'Verificación de operación {numero_operacion}', style=self.list_number)
                    time.sleep(3)
                    self.document.add_picture(img_verificacion_operacion, width=Inches(6.73))
                    self.action_bot.presionar_boton(self.btn_aceptar_ventana)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_aceptar_ventana)
                    print(f"Fila {self.e} ingresada con exito")
                    time.sleep(3)
                    self.cerrar_sesion()
                    self.e+=1
                else:
                    print("Se ha ingresado un usuario o contraseña incorrecta, se continuará con la segunda fila")
                    time.sleep(5)
                    img_error_usuario = f"./Screenshots/{self.e}_MDP_ERROR_USUARIO_{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
                    screenshot1 = pyautogui.screenshot()
                    screenshot1.save(img_error_usuario)
                    #Inserción en Word
                    self.document.add_paragraph(f'Error al ingresar al usuario de la fila {self.e}', style=self.list_number)
                    self.document.add_picture(img_error_usuario, width=Inches(6.73))
                    time.sleep(5)
                    self.action_bot.presionar_boton(self.btn_error_usuario)
                    self.e+=1
            except Exception as e:
                print(f"Error en fila {self.e}: {e}")
                # Continúa con la siguiente fila del Excel
                self.e += 1
                time.sleep(5)
                self.cerrar_sesion()
                continue    

            except TimeoutException as e:
                print(f"Se ha acabado el tiempo de espera, Exception :{e}")
            except NoSuchElementException as e:
                print(f"NoSuchElement {e}")
            except ElementNotInteractableException as e:
                print(f"ElementNotIteractable {e}")
            except StaleElementReferenceException as e:
                print(f"StaleElement {e}")
            except ElementClickInterceptedException as e:
                print(f"ElementClickInter {e}")
           

    def verificar_con_verificador(self,numero_operacion: str):
        print(f"Verificando operación {numero_operacion}...")
        self.action_bot.presionar_boton(self.btn_menu_transacciones)
        self.action_bot.presionar_boton(self.btn_menu_verificar)
        self.action_bot.presionar_boton(self.btn_numero_operacion)
        self.action_bot.ingresar_texto(self.txt_numero_operacion, self.detect_type(numero_operacion) )
        self.action_bot.presionar_boton(self.btn_numero_operacion_buscar)
        self.action_bot.presionar_boton(self.casilla_numero_operacion)
        time.sleep(1)
        self.action_bot.presionar_boton(self.btn_aprobar_verificar)

    def consultar_operacion(self,numero_operacion: str):
        print(f"Consultando operación {numero_operacion}...")
        self.action_bot.presionar_boton(self.btn_menu_gestion)
        time.sleep(2)
        self.action_bot.presionar_boton(self.btn_menu_consulta)
        time.sleep(2)
        self.action_bot.presionar_boton(self.btn_numero_operacion)
        time.sleep(1)
        self.action_bot.ingresar_texto(self.txt_numero_operacion, self.detect_type(numero_operacion) )
        self.action_bot.presionar_boton(self.btn_numero_operacion_buscar)

    
    def ingresar_mt202(self,i): #i: xlwings.main.Range
        print("Se esta ingresando una operación MT202...")
        self.slc_pago_producto(i)
        try:
            print("Verificando rut...")
            self.action_bot.ingresar_texto(self.txt_rut_ingresar, self.detect_type(i[4].value))
            self.action_bot.presionar_boton(self.txt_nombre_ingresar)
            if self.validarRut(self.detect_type(i[4].value)) is True:
                if "." not in self.detect_type(i[4].value):
                    time.sleep(3)
                    print("Rut Correcto")
                    self.monto_y_continuar(i)
                    print("Por confirmar operación...")
                else:
                    print("El rut se debe ingresar sin puntos. Formato 12345678-K")
                    raise Exception("El rut se debe ingresar sin puntos. Formato 12345678-K") 
            else:
                print("Se ha ingresado un rut invalido")
                raise Exception("Se ha ingresado un rut invalido")

        except Exception as e:
            time.sleep(5)
            img_error_usuario = f"./Screenshots/{self.e}_MDP_ERROR_USUARIO_{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
            screenshot5 = pyautogui.screenshot()
            screenshot5.save(img_error_usuario)
            #Inserción en Word
            self.document.add_paragraph(f'Error al ingresar rut de la fila {self.e}', style=self.list_number)
            self.document.add_picture(img_error_usuario, width=Inches(6.73))
            time.sleep(5)
            self.action_bot.presionar_boton(self.btn_aceptar_ventana)
            raise
        

    def ingresar_mt103(self,i): #i: xlwings.main.Range
        print("Se esta ingresando una operación MT103...")
        self.slc_pago_producto(i)
        try:
            print("Verificando rut...")
            self.action_bot.ingresar_texto(self.txt_rut_ingresar, self.detect_type(i[4].value))
            self.action_bot.presionar_boton(self.txt_nombre_ingresar)
            if self.validarRut(self.detect_type(i[4].value)) is True:
                if "." not in self.detect_type(i[4].value):
                    print("Rut Correcto")
                    time.sleep(3)
                    self.action_bot.ingresar_texto(self.txt_nombre_ingresar, i[6].value)
                    time.sleep(3)
                    self.monto_y_continuar(i)
                    self.action_bot.select_text(self.slc_banco, self.detect_type(i[8].value))
                    time.sleep(3)
                    print("Se esta ingresando el beneficiario...")
                    self.action_bot.presionar_boton(self.btn_beneficiario_ingresar)
                    time.sleep(3)
                    self.action_bot.ingresar_texto(self.txt_cuenta_ingresar, self.detect_type(i[9].value))
                    self.action_bot.presionar_boton(self.btn_grabar_beneficiario)
                    time.sleep(3)
                    self.action_bot.presionar_boton(self.btn_aceptar_ventana)
                    time.sleep(5)
                else:
                    print("Rut incorrecto")
                    raise Exception("El rut se debe ingresar sin puntos. Formato 12345678-K") 
            else:
                raise Exception("Se ha ingresado un rut invalido")

        except Exception as e:
            time.sleep(5)
            img_error_usuario = f"./Screenshots/{self.e}_MDP_ERROR_USUARIO_{datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.png"
            screenshot5 = pyautogui.screenshot()
            screenshot5.save(img_error_usuario)
            #Inserción en Word
            self.document.add_paragraph(f'Error al ingresar rut de la fila {self.e}', style=self.list_number)
            self.document.add_picture(img_error_usuario, width=Inches(6.73))
            time.sleep(5)
            self.action_bot.presionar_boton(self.btn_aceptar_ventana)
            raise
        

    def validarRut(self,rut: str):
        rut = rut.upper()
        rut = rut.replace("-","")
        #rut = rut.replace(".","")
        aux = rut[:-1]
        dv = rut[-1:]
    
        revertido = map(int, reversed(str(aux)))
        factors = cycle(range(2,8))
        s = sum(d * f for d, f in zip(revertido,factors))
        res = (-s)%11
    
        if str(res) == dv:
            return True
        elif dv=="K" and res==10:
            return True
        else:
            return False

    def slc_pago_producto(self,i): #i: xlwings.main.Range
        print("Seleccionando forma de pago...")
        self.action_bot.select_text(self.slc_forma_pago, i[2].value)
        print("Seleccionando producto...")
        time.sleep(15)
        self.action_bot.select_text(self.slc_producto, self.detect_type(i[3].value))
        time.sleep(3)

    def monto_y_continuar(self,i): #i: xlwings.main.Range
        self.action_bot.ingresar_texto(self.txt_monto_ingresar, self.detect_type(i[7].value))
        self.action.send_keys(Keys.ENTER).perform()
        time.sleep(5)
        self.action_bot.presionar_boton(self.btn_continuar_ingresar)
        time.sleep(5)
    
    def iniciar_sesion_verificador(self,i): #i: xlwings.main.Range
        username_verificador = self.detect_type(i[11].value)
        password_verificador = self.detect_type(i[12].value)
        self.action_bot.ingresar_texto(self.txt_username, username_verificador)
        self.action_bot.ingresar_texto(self.txt_password, password_verificador)
        self.action.send_keys(Keys.ENTER).perform()


    def iniciar_sesion_ingresador(self,i): #i: xlwings.main.Range
        username_ingresador = self.detect_type(i[0].value)
        password_ingresador = self.detect_type(i[1].value)
        self.action_bot.ingresar_texto(self.txt_username, username_ingresador)
        self.action_bot.ingresar_texto(self.txt_password, password_ingresador)
        self.action.send_keys(Keys.ENTER).perform()

    def cerrar_sesion(self):
        self.action_bot.presionar_boton(self.btn_cerrar_sesion)
        time.sleep(2)
        self.action_bot.presionar_boton(self.btn_aceptar_cerrar_sesion)
        time.sleep(2)

    def detect_type(self,valorr):         #No se declara type date de 'valorr': str or float ya que está función recibe String y Float
        if isinstance(valorr,float):
            return str(int(valorr))
        else:
            return valorr

    def vaciar_screens(self):
        directorio = "./Screenshots"
        # Obtener la lista de archivos en el directorio
        archivos = os.listdir(directorio)
        for archivo in archivos:
            # Verificar si el archivo es una imagen (por ejemplo, PNG)
            if archivo.endswith(".png"):
                ruta_archivo = os.path.join(directorio, archivo) # Construir la ruta completa del archivo
                send2trash(ruta_archivo) # Eliminar el archivo

    def terminar_ejecucion(self):
        self.workbook.save()
        self.workbook.close()
        self.app.quit()
        self.document.save(f'./Reportes/Reporte {datetime.now().strftime("%m-%d-%Y_%H;%M;%S")}.docx')
        self.vaciar_screens()
        print("Reporte guardado con exito")
        input("Presiona Enter para detener el bot... \nADVERTENCIA: AL DETENER EL BOT TAMBIEN SE CIERRA LA PESTANA ACTUAL, \nSE RECOMIENDA SOLO DETENER CUANDO YA HAZ TERMINADO EL USO DE ESTA PESTANA")
        self.driver.quit()
        sys.exit()

if __name__ == '__main__':
    init_bot = Motor(archivo_xlsx = 'Automatizacion.xlsx')
    try:
        init_bot.ejecucion_motor()
    except ElementClickInterceptedException as e:
        print(e)
    except TimeoutException as ep:
        print(ep)
    except Exception as e:
        print(e)
    finally:
        init_bot.terminar_ejecucion()    
